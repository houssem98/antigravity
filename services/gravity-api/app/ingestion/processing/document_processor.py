"""
Gravity Search — Document Processor
Converts raw bytes (HTML, PDF, DOCX, plain text) into clean plain text.
Preserves document structure (headings, paragraphs) for downstream section detection.

Enhanced with table extraction per Table Reasoning papers:
  - HTML tables → ParsedTable objects (most SEC filings are HTML)
  - PDF tables  → ParsedTable via pdfplumber
  Tables are stored separately from text to preserve row/column structure.
"""

import structlog
from dataclasses import dataclass, field

from app.ingestion.processing.table_parser import FinancialTableParser, ParsedTable

logger = structlog.get_logger()

# Shared table parser instance
_table_parser = FinancialTableParser()


@dataclass
class ProcessedDocument:
    """Result of processing a raw document."""
    text: str              # Clean plain text
    title: str = ""        # Extracted document title
    page_count: int = 0    # Number of pages (PDF only)
    language: str = "en"   # Detected language code
    raw_html: str = ""     # Original HTML (for debugging)
    tables: list = field(default_factory=list)    # list[ParsedTable] — structured tables
    xbrl_facts: list = field(default_factory=list)  # list[XBRLFact] — machine-readable facts


class DocumentProcessor:
    """
    Normalizes documents from various formats into clean plain text.

    Supported formats:
    - HTML (.html, text/html)
    - PDF (.pdf, application/pdf)
    - DOCX (.docx)
    - Plain text (.txt, text/plain)
    """

    async def process(
        self,
        content: bytes,
        content_type: str = "",
        filename: str = "",
    ) -> ProcessedDocument:
        """
        Main entry point. Detects format and dispatches to the correct processor.

        Args:
            content: Raw document bytes
            content_type: MIME type from HTTP header
            filename: Original filename (used for format detection)

        Returns:
            ProcessedDocument with clean text ready for chunking.
        """
        ct = (content_type or "").lower()
        fn = (filename or "").lower()

        if "html" in ct or fn.endswith((".html", ".htm")):
            return self._process_html(content)
        elif "pdf" in ct or fn.endswith(".pdf"):
            return self._process_pdf(content)
        elif fn.endswith(".docx"):
            return self._process_docx(content)
        else:
            # Default: treat as plain text
            return self._process_text(content)

    def _process_html(self, content: bytes) -> ProcessedDocument:
        """Strip HTML tags and extract clean text + structured tables."""
        try:
            from bs4 import BeautifulSoup
            raw_html = content.decode("utf-8", errors="replace")
            soup = BeautifulSoup(raw_html, "lxml")

            # Extract title
            title_tag = soup.find("title")
            title = title_tag.get_text(strip=True) if title_tag else ""

            # ── Extract structured tables BEFORE removing elements ────────
            tables = _table_parser.extract_html_tables(raw_html)

            # ── Extract XBRL facts from iXBRL-embedded HTML ───────────────
            xbrl_facts: list = []
            try:
                from app.ingestion.processing.xbrl_extractor import XBRLExtractor, DerivedMetricsCalculator
                _xbrl = XBRLExtractor()
                xbrl_facts = _xbrl.extract_from_html(raw_html)
                # Add derived metrics (margins, ratios) to improve coverage
                derived = DerivedMetricsCalculator.compute(xbrl_facts)
                xbrl_facts.extend(derived)
            except Exception as _xe:
                logger.debug("xbrl_extraction_skipped", error=str(_xe))

            # Remove noise elements
            for tag in soup(["script", "style", "nav", "footer", "header",
                              "aside", "noscript", "form", "button"]):
                tag.decompose()

            # Preserve headings and paragraphs as double-newline separated text
            text_parts = []
            for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p",
                                           "li", "td", "th", "div"]):
                t = element.get_text(separator=" ", strip=True)
                if t and len(t) > 10:  # Skip very short fragments
                    text_parts.append(t)

            text = "\n\n".join(text_parts)
            if not text:
                # Fallback: get all text
                text = soup.get_text(separator="\n", strip=True)

            logger.info("html_processed", title=title[:50], chars=len(text), tables=len(tables), xbrl_facts=len(xbrl_facts))
            return ProcessedDocument(
                text=text, title=title, raw_html=raw_html[:10000], tables=tables, xbrl_facts=xbrl_facts,
            )

        except Exception as e:
            logger.warning("html_process_failed", error=str(e))
            # Fallback: strip HTML with regex
            import re
            text = re.sub(r"<[^>]+>", " ", content.decode("utf-8", errors="replace"))
            text = re.sub(r"\s+", " ", text).strip()
            return ProcessedDocument(text=text)

    def _process_pdf(self, content: bytes) -> ProcessedDocument:
        """Extract text + structured tables from PDF."""
        import io

        # ── Extract structured tables via pdfplumber ──────────────────
        tables = _table_parser.extract_pdf_tables(content)

        # ── Extract text via pdfminer.six (better text extraction) ───
        text = ""
        page_count = 0
        try:
            from pdfminer.high_level import extract_text
            text = extract_text(io.BytesIO(content))
            if text and text.strip():
                page_count = self._count_pdf_pages(content)
                text = text.strip()
        except ImportError:
            pass
        except Exception as e:
            logger.warning("pdfminer_failed", error=str(e))

        # Fallback: pypdf
        if not text:
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(content))
                pages = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
                text = "\n\n".join(pages)
                page_count = len(reader.pages)
            except ImportError:
                pass
            except Exception as e:
                logger.warning("pypdf_failed", error=str(e))

        if not text:
            logger.warning("pdf_all_parsers_failed", size=len(content))
            text = content.decode("utf-8", errors="replace")

        logger.info("pdf_processed", chars=len(text), pages=page_count, tables=len(tables))
        return ProcessedDocument(text=text, page_count=page_count, tables=tables)

    def _count_pdf_pages(self, content: bytes) -> int:
        """Count PDF pages without extracting text."""
        try:
            import pypdf, io
            return len(pypdf.PdfReader(io.BytesIO(content)).pages)
        except Exception:
            return 0

    def _process_docx(self, content: bytes) -> ProcessedDocument:
        """Extract text from DOCX using python-docx."""
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            # Extract title from core properties
            title = ""
            try:
                title = doc.core_properties.title or ""
            except Exception:
                pass
            logger.info("docx_processed", title=title[:50], paragraphs=len(paragraphs))
            return ProcessedDocument(text=text, title=title)
        except Exception as e:
            logger.warning("docx_process_failed", error=str(e))
            return ProcessedDocument(text=content.decode("utf-8", errors="replace"))

    def _process_text(self, content: bytes) -> ProcessedDocument:
        """Decode plain text, normalize whitespace."""
        try:
            text = content.decode("utf-8", errors="replace")
        except Exception:
            text = content.decode("latin-1", errors="replace")

        # Normalize: collapse excessive blank lines
        import re
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        return ProcessedDocument(text=text)
